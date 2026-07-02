# -*- coding: utf-8 -*-
"""Bibliotecă reutilizabilă pentru portofolii academice .docx (stil RO, Times New Roman 12 / 1.5).
Adaptează `COVER` și funcția `build_body` per proiect, sau importă helperele într-un content.py separat.

    from docx_template import *
    def build_body(doc):
        cuprins(doc, ["I. ...", "II. ...", "Bibliografie"])
        h1(doc, "I. Titlu"); para(doc, "text...", first_indent=1.25)
        swot_table(doc, tari=[...], slabe=[...], oport=[...], amenint=[...])
        bibliography(doc, ["Autor, A. (2024). Titlu. Editura."])
    make_document(COVER, build_body, "Portofoliu.docx", logo_path="logo.png")

Necesită: pip install python-docx ; logo PNG opțional.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Times New Roman"

# Exemplu de metadate pentru copertă (înlocuiește per proiect):
COVER = {
    "universitate": "NUMELE UNIVERSITĂȚII",
    "facultate": "Facultatea …",
    "departament": "Departamentul …",
    "extra_linii": ["Studii universitare de … (licență / masterat)",
                    "Programul de studiu: …",
                    "Anul universitar 20XX–20XX"],
    "titlu": "PROIECT / PORTOFOLIU DE EXAMEN",
    "subtitlu": "Titlul temei / disciplinei",
    "titular_label": "Titular de curs:",
    "titular": "Grad. dr. Nume Prenume",
    "student_label": "Student(ă):",
    "student": ["Nume Prenume", "An / Program de studiu"],
    "oras_an": "Oraș, 20XX",
    "logo_width_cm": 7.0,
}


# ----------------------------- helpers de bază -----------------------------

def _set_run(run, size=12, bold=False, italic=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_base_styles(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_margins(section, top=2.5, bottom=2.5, left=2.5, right=2.5):
    section.top_margin, section.bottom_margin = Cm(top), Cm(bottom)
    section.left_margin, section.right_margin = Cm(left), Cm(right)


_ALIGN = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
          "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}


def para(doc, text="", size=12, bold=False, italic=False, align="justify",
         space_after=6, space_before=0, line=1.5, color=None, first_indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = _ALIGN[align]
    pf.space_after, pf.space_before = Pt(space_after), Pt(space_before)
    if line == 1.5:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    else:
        pf.line_spacing = line
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if text:
        _set_run(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(14), Pt(8)
    p.paragraph_format.keep_with_next = True
    _set_run(p.add_run(text), size=14, bold=True)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(10), Pt(4)
    p.paragraph_format.keep_with_next = True
    _set_run(p.add_run(text), size=12.5, bold=True)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(6), Pt(2)
    p.paragraph_format.keep_with_next = True
    _set_run(p.add_run(text), size=12, bold=True, italic=True)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_run(p.add_run(text), size=12)
    return p


def blockquote(doc, text, source):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent, pf.right_indent = Cm(1.25), Cm(1.0)
    pf.space_before, pf.space_after = Pt(6), Pt(6)
    pf.line_spacing = 1.0
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_run(p.add_run("„" + text + "” "), size=11, italic=True)
    _set_run(p.add_run(source), size=11)
    return p


def cuprins(doc, items):
    h1(doc, "Cuprins")
    for t in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        _set_run(p.add_run(t), size=12)
    doc.add_page_break()


def swot_table(doc, tari, slabe, oport, amenint):
    t = doc.add_table(rows=2, cols=2)
    t.style, t.alignment = "Table Grid", 1
    cells = [(0, 0, "Puncte tari (Strengths)", tari), (0, 1, "Puncte slabe (Weaknesses)", slabe),
             (1, 0, "Oportunități (Opportunities)", oport), (1, 1, "Amenințări (Threats)", amenint)]
    for r, c, title, items in cells:
        cell = t.cell(r, c)
        cell.width = Cm(8)
        _set_run(cell.paragraphs[0].add_run(title), size=11, bold=True)
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)
        for it in items:
            ip = cell.add_paragraph()
            ip.paragraph_format.space_after, ip.paragraph_format.line_spacing = Pt(2), 1.0
            _set_run(ip.add_run("• " + it), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def bibliography(doc, refs):
    h1(doc, "Bibliografie")
    for r in sorted(refs, key=lambda s: s.lower()):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent, pf.first_line_indent = Cm(1.0), Cm(-1.0)  # hanging indent
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_run(p.add_run(r), size=12)


def page_number_footer(section):
    section.different_first_page_header_footer = True  # fără număr pe copertă
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for tag, attr, val, txt in [("begin", None, None, None), (None, "instr", None, "PAGE"), ("end", None, None, None)]:
        if tag:
            el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), tag)
        else:
            el = OxmlElement("w:instrText"); el.set(qn("xml:space"), "preserve"); el.text = txt
        run._r.append(el)
    _set_run(run, size=10)


def build_cover(doc, meta, logo_path=None):
    set_margins(doc.sections[0])
    if logo_path:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(6), Pt(10)
        p.add_run().add_picture(logo_path, width=Cm(meta.get("logo_width_cm", 7.0)))
    para(doc, meta["universitate"], size=13, bold=True, align="center", space_after=2, line=1.0)
    para(doc, meta["facultate"], align="center", space_after=2, line=1.0)
    para(doc, meta["departament"], align="center", space_after=2, line=1.0)
    for ln in meta.get("extra_linii", []):
        para(doc, ln, align="center", space_after=2, line=1.0)
    for _ in range(4):
        para(doc, "", line=1.0, space_after=0)
    para(doc, meta["titlu"], size=22, bold=True, align="center", space_after=6, line=1.0)
    para(doc, meta["subtitlu"], size=16, italic=True, align="center", space_after=4, line=1.0)
    for _ in range(5):
        para(doc, "", line=1.0, space_after=0)
    para(doc, meta["titular_label"], bold=True, align="left", space_after=0, line=1.0)
    para(doc, meta["titular"], align="left", space_after=10, line=1.0)
    para(doc, meta["student_label"], bold=True, align="left", space_after=0, line=1.0)
    for ln in meta["student"]:
        para(doc, ln, align="left", space_after=0, line=1.0)
    for _ in range(6):
        para(doc, "", line=1.0, space_after=0)
    para(doc, meta["oras_an"], align="center", space_after=0, line=1.0)
    doc.add_page_break()


def make_document(meta, build_body, out_path, logo_path=None):
    doc = Document()
    set_base_styles(doc)
    build_cover(doc, meta, logo_path)
    build_body(doc)
    page_number_footer(doc.sections[0])
    doc.save(out_path)
    print("Saved", out_path)
    return out_path
