# -*- coding: utf-8 -*-
"""Construiește un SUPORT DE ÎNVĂȚAT (.docx) clar și scanabil din module distilate.
Două moduri: 'condensat' (sinteză rapidă) și 'detaliat' (explicativ, amplu).
Autonom (nu depinde de docx_template).

Utilizare:
    python3 study_builder.py <condensat|detaliat> module.json out.docx "TITLU" "Subtitlu"

module.json = {"modules": [ ... ]} unde fiecare modul respectă schema din SKILL.md
(chapter_no, chapter_title, pe_scurt/introducere, concepte, clasificari, procese,
 formule, exemple/aplicatii, mnemonice, confuzii, rezumat, intrebari).
"""
import sys, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Calibri"
NAVY = RGBColor(0x1F, 0x3A, 0x5F); TEAL = RGBColor(0x1E, 0x6E, 0x6E); INK = RGBColor(0x22, 0x22, 0x22)
GREEN = RGBColor(0x2E, 0x7D, 0x4F); AMBER = RGBColor(0x9A, 0x6A, 0x00); MUT = RGBColor(0x66, 0x66, 0x66)


def _run(r, size=11, bold=False, italic=False, color=INK, font=FONT):
    r.font.name = font; r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color; r._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def _shade(p, hexfill):
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hexfill); p._p.get_or_add_pPr().append(sh)


def base_styles(doc):
    st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(4); st.paragraph_format.line_spacing = 1.12


def P(doc, text="", size=11, bold=False, italic=False, color=INK, after=4, align="left"):
    p = doc.add_paragraph()
    p.alignment = {"left": 0, "center": 1, "just": 3}[align]
    p.paragraph_format.space_after = Pt(after)
    if text: _run(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def H1(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True; _shade(p, "1F3A5F")
    _run(p.add_run("  " + text), size=16, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    return p


def H2(doc, text, color=TEAL):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    _run(p.add_run(text), size=12.5, bold=True, color=color)
    pPr = p._p.get_or_add_pPr(); pb = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), "BBBBBB"); pb.append(bot); pPr.append(pb)
    return p


def H3(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    _run(p.add_run(text), size=11.5, bold=True, color=NAVY)
    return p


def box(doc, label, text, fill="EAF1F8"):
    p = doc.add_paragraph(); _shade(p, fill)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.2); p.paragraph_format.right_indent = Cm(0.2)
    if label: _run(p.add_run(label + " "), size=11, bold=True, color=NAVY)
    _run(p.add_run(text), size=11, color=INK)


def bullet(doc, text=None, runs=None, marker="•", color=INK, size=11, indent=0.6):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(indent); p.paragraph_format.space_after = Pt(2)
    if marker: _run(p.add_run(marker + "  "), size=size, bold=True, color=color)
    if runs:
        for t, opt in runs:
            _run(p.add_run(t), size=size, bold=opt.get("b", False), italic=opt.get("i", False), color=opt.get("c", INK))
    else:
        _run(p.add_run(text), size=size, color=INK)


def deffist(doc, term, definition):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.space_after = Pt(2)
    _run(p.add_run(term + " — "), size=11, bold=True, color=NAVY); _run(p.add_run(definition), size=11)


def explic(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12; p.alignment = 3
    _run(p.add_run(text), size=11, color=INK)


def _chapter_label(m):
    no = str(m.get("chapter_no", "")); title = m.get("chapter_title", "")
    return f"Capitolul {no}. {title}" if (no.replace('/', '').isdigit() or '/' in no) else title


def module_condensed(doc, m):
    H1(doc, _chapter_label(m))
    if m.get("pe_scurt"): box(doc, "Pe scurt:", m["pe_scurt"])
    if m.get("concepte"):
        H2(doc, "Concepte-cheie")
        for c in m["concepte"]: deffist(doc, c.get("termen", ""), c.get("definitie", ""))
    if m.get("clasificari"):
        H2(doc, "Clasificări și tipologii")
        for cl in m["clasificari"]:
            H3(doc, cl.get("titlu", ""))
            for el in cl.get("elemente", []):
                bullet(doc, el if isinstance(el, str) else (el.get("nume", "") + " — " + el.get("descriere", "")))
    if m.get("formule"):
        H2(doc, "Formule")
        for f in m["formule"]:
            runs = [(f.get("nume", "") + ": ", {"b": True, "c": NAVY}), (f.get("formula", ""), {"b": True, "c": TEAL})]
            if f.get("explicatie"): runs.append(("  — " + f["explicatie"], {"i": True, "c": MUT}))
            bullet(doc, runs=runs, marker="f", color=TEAL)
    for key, head in [("exemple", "Exemple"), ("aplicatii", "Exemple")]:
        if m.get(key):
            H2(doc, head)
            for e in m[key]: bullet(doc, e)
            break
    if m.get("mnemonice"):
        H2(doc, "De memorat (trucuri)", color=GREEN)
        for mn in m["mnemonice"]: bullet(doc, mn, marker="+", color=GREEN)
    if m.get("confuzii"):
        H2(doc, "Atenție — confuzii frecvente", color=AMBER)
        for cf in m["confuzii"]:
            bullet(doc, cf if isinstance(cf, str) else (cf.get("confuzie", "") + " → " + cf.get("clarificare", "")),
                   marker="!", color=AMBER)
    if m.get("intrebari"):
        H2(doc, "Test rapid (autoevaluare)")
        for i, qa in enumerate(m["intrebari"], 1):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.space_after = Pt(1)
            _run(p.add_run(f"{i}. {qa.get('q','')}"), size=11, bold=True)
            pa = doc.add_paragraph(); pa.paragraph_format.left_indent = Cm(0.8); pa.paragraph_format.space_after = Pt(5)
            _run(pa.add_run("R: "), size=10.5, bold=True, color=GREEN); _run(pa.add_run(qa.get("a", "")), size=10.5, italic=True, color=MUT)
    doc.add_page_break()


def module_detailed(doc, m):
    H1(doc, _chapter_label(m))
    if m.get("introducere"):
        for para in [s for s in m["introducere"].split("\n") if s.strip()]: explic(doc, para.strip())
    elif m.get("pe_scurt"):
        box(doc, "Pe scurt:", m["pe_scurt"])
    if m.get("concepte"):
        H2(doc, "Concepte explicate")
        for c in m["concepte"]:
            deffist(doc, c.get("termen", ""), c.get("definitie", ""))
            if c.get("explicatie"): explic(doc, c["explicatie"])
    if m.get("clasificari"):
        H2(doc, "Clasificări și tipologii (explicate)")
        for cl in m["clasificari"]:
            H3(doc, cl.get("titlu", ""))
            if cl.get("intro"): explic(doc, cl["intro"])
            for el in cl.get("elemente", []):
                if isinstance(el, dict):
                    bullet(doc, runs=[(el.get("nume", "") + " — ", {"b": True, "c": NAVY}), (el.get("descriere", ""), {})])
                else:
                    bullet(doc, el)
    if m.get("procese"):
        H2(doc, "Procese și etape")
        for pr in m["procese"]:
            H3(doc, pr.get("titlu", ""))
            for i, st in enumerate(pr.get("pasi", []), 1):
                bullet(doc, runs=[(f"{i}. " + st.get("pas", "") + " — ", {"b": True, "c": TEAL}), (st.get("descriere", ""), {})], marker="", indent=0.5)
    if m.get("formule"):
        H2(doc, "Formule (cu exemplu)")
        for f in m["formule"]:
            bullet(doc, runs=[(f.get("nume", "") + ": ", {"b": True, "c": NAVY}), (f.get("formula", ""), {"b": True, "c": TEAL})], marker="f", color=TEAL)
            if f.get("explicatie"): explic(doc, f["explicatie"])
            if f.get("exemplu"):
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.space_after = Pt(6)
                _run(p.add_run("Exemplu: "), size=10.5, bold=True, color=GREEN); _run(p.add_run(f["exemplu"]), size=10.5, italic=True, color=MUT)
    if m.get("aplicatii") or m.get("exemple"):
        H2(doc, "Aplicații și exemple")
        for a in (m.get("aplicatii") or m.get("exemple")): bullet(doc, a)
    if m.get("confuzii"):
        H2(doc, "Confuzii frecvente clarificate", color=AMBER)
        for cf in m["confuzii"]:
            if isinstance(cf, dict):
                bullet(doc, runs=[(cf.get("confuzie", "") + " → ", {"b": True, "c": AMBER}), (cf.get("clarificare", ""), {})], marker="!", color=AMBER)
            else:
                bullet(doc, cf, marker="!", color=AMBER)
    if m.get("rezumat"):
        H2(doc, "Rezumat"); box(doc, "", m["rezumat"])
    if m.get("intrebari"):
        H2(doc, "Întrebări de verificare (cu răspuns)")
        for i, qa in enumerate(m["intrebari"], 1):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.space_after = Pt(1)
            _run(p.add_run(f"{i}. {qa.get('q','')}"), size=11, bold=True)
            pa = doc.add_paragraph(); pa.paragraph_format.left_indent = Cm(0.8); pa.paragraph_format.space_after = Pt(6); pa.paragraph_format.line_spacing = 1.12; pa.alignment = 3
            _run(pa.add_run("R: "), size=10.5, bold=True, color=GREEN); _run(pa.add_run(qa.get("a", "")), size=10.5)
    doc.add_page_break()


def build(mode, modules_json, out_path, title="SUPORT DE ÎNVĂȚAT", subtitle=""):
    data = json.load(open(modules_json)); mods = data.get("modules", data)
    def k(m):
        try: return int(str(m.get("chapter_no", "99")))
        except: return 99
    mods = sorted(mods, key=k)
    doc = Document(); base_styles(doc)
    sec = doc.sections[0]; sec.top_margin = sec.bottom_margin = Cm(2.0); sec.left_margin = sec.right_margin = Cm(2.2)
    # cover
    for _ in range(4): P(doc, "", after=0)
    P(doc, title, size=30, bold=True, color=NAVY, align="center", after=6)
    if subtitle: P(doc, subtitle, size=15, italic=True, color=MUT, align="center", after=18)
    P(doc, ("Versiune " + ("condensată" if mode == "condensat" else "detaliată")), size=13, color=INK, align="center")
    doc.add_page_break()
    # cuprins
    H1(doc, "Cuprins")
    for m in mods:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5); _run(p.add_run(_chapter_label(m)), size=12)
    doc.add_page_break()
    render = module_condensed if mode == "condensat" else module_detailed
    for m in mods: render(doc, m)
    doc.save(out_path); print("Saved", out_path, "| module:", len(mods))


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print('Utilizare: python3 study_builder.py <condensat|detaliat> module.json out.docx "TITLU" "Subtitlu"'); sys.exit(1)
    build(sys.argv[1], sys.argv[2], sys.argv[3],
          sys.argv[4] if len(sys.argv) > 4 else "SUPORT DE ÎNVĂȚAT",
          sys.argv[5] if len(sys.argv) > 5 else "")
