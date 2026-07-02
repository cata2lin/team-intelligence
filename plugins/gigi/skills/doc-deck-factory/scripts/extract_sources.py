# -*- coding: utf-8 -*-
"""Extrage text (inclusiv tabele) din TOATE documentele dintr-un folder,
pentru a le da apoi agenților de sinteză. Convertește .doc vechi -> .docx (textutil, macOS).

Utilizare:
    python3 extract_sources.py "/cale/folder_sursa" [/cale/out_dir]
Implicit out_dir = /tmp/dd_src . Preferă .docx când există și .doc și .docx cu același nume.
Necesită: python-docx ; (macOS: textutil pentru .doc). Pe alt OS, convertește .doc manual întâi.
"""
import sys, os, glob, subprocess
from docx import Document


def extract_docx(path):
    d = Document(path)
    lines = []
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for tb in d.tables:
        for r in tb.rows:
            cells = [c.text.strip().replace("\n", " ") for c in r.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def main(src_dir, out_dir="/tmp/dd_src"):
    os.makedirs(out_dir, exist_ok=True)
    conv_dir = os.path.join(out_dir, "_converted")
    os.makedirs(conv_dir, exist_ok=True)
    # 1) convert .doc -> .docx (macOS textutil)
    for f in glob.glob(os.path.join(src_dir, "*.doc")):
        out = os.path.join(conv_dir, os.path.splitext(os.path.basename(f))[0] + ".docx")
        try:
            subprocess.run(["textutil", "-convert", "docx", f, "-output", out], check=True,
                           capture_output=True)
        except Exception as e:
            print("WARN nu am putut converti", f, e)
    # 2) extract all docx (originals + converted); dedupe by base name (prefer original .docx)
    seen = {}
    for path in sorted(glob.glob(os.path.join(src_dir, "*.docx"))) + sorted(glob.glob(os.path.join(conv_dir, "*.docx"))):
        base = os.path.splitext(os.path.basename(path))[0].strip().lower()
        if base in seen:
            continue  # original .docx wins over converted .doc
        seen[base] = path
    for base, path in seen.items():
        try:
            txt = extract_docx(path)
            name = os.path.splitext(os.path.basename(path))[0]
            open(os.path.join(out_dir, name + ".txt"), "w").write(txt)
            print(f"{name}: {len(txt.split())} cuvinte -> {out_dir}/{name}.txt")
        except Exception as e:
            print("ERR", path, e)
    print(f"\nGata. {len(seen)} fișiere text în {out_dir}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print('Utilizare: python3 extract_sources.py "/folder_sursa" [/out_dir]'); sys.exit(1)
    main(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else "/tmp/dd_src")
